import logging
import sys
import os
import traceback
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import yaml
import re
import argparse

# --- Default Configuration Constants ---
DEFAULT_INPUT_FILE = 'data/input/input.md'
DEFAULT_OUTPUT_FILE = 'data/output/output.docx'
LOG_FILE = 'logs/generate_ieee_format.log'

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


def read_markdown_file(filepath):
    """Read the markdown file with proper error handling."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info(f"Successfully read {filepath}")
        return content
    except FileNotFoundError:
        logger.error(f"Error: {filepath} not found. Please ensure the file exists in the current directory.")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Unexpected error reading markdown file: {str(e)}")
        sys.exit(1)


def parse_frontmatter(content):
    """Extract YAML frontmatter from the markdown content."""
    try:
        match = re.match(r'^---\n(.*?)\n---\n(.*)', content, re.DOTALL | re.MULTILINE)
        if match:
            frontmatter = yaml.safe_load(match.group(1))
            body_content = match.group(2).strip()
            logger.info("Successfully parsed YAML frontmatter metadata.")
            return frontmatter, body_content
        else:
            logger.warning("No YAML frontmatter found. Using default placeholder values.")
            return {}, content
    except yaml.YAMLError as exc:
        logger.error(f"Error parsing YAML frontmatter: {exc}")
        return {}, content
    except Exception as e:
        logger.error(f"Unexpected error parsing metadata: {str(e)}")
        return {}, content


def setup_page_margins(doc):
    """Set the margins for the main document."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)


def add_title(doc, frontmatter):
    """Add the main paper title."""
    title_text = frontmatter.get('title', 'Unknown Title')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    title.paragraph_format.space_after = Pt(18)


def add_authors(doc, frontmatter):
    """Add author dynamically depending on YAML frontmatter."""
    authors = frontmatter.get('authors', [])
    num_authors = len(authors) if authors else 1
    author_table = doc.add_table(rows=1, cols=max(1, num_authors))
    author_table.autofit = False
    author_table.allow_autofit = False

    # Remove borders
    for row in author_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    for idx, author in enumerate(authors):
        if idx >= len(author_table.rows[0].cells):
            break
        cell = author_table.rows[0].cells[idx]
        cell.width = Inches(7.25 / num_authors)
        
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(author.get('name', ''))
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p1.paragraph_format.space_after = Pt(3)
        
        if 'role' in author:
            p_role = cell.add_paragraph()
            p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_role.add_run(author['role'])
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p_role.paragraph_format.space_after = Pt(3)
            
        if 'department' in author:
            p_dept = cell.add_paragraph()
            p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_dept.add_run(author['department'])
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p_dept.paragraph_format.space_after = Pt(3)
            
        if 'organization' in author:
            p_org = cell.add_paragraph()
            p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_org.add_run(author['organization'])
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p_org.paragraph_format.space_after = Pt(3)
            
        if 'email' in author:
            p_email = cell.add_paragraph()
            p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_email.add_run(author['email'])
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p_email.paragraph_format.space_after = Pt(0)

    # Add spacing after author table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def set_two_column_layout(doc):
    """Set layout to double columns for main text content."""
    new_section = doc.add_section()
    new_section.start_type = 0  # Continuous section break
    new_section.page_height = Inches(11)
    new_section.page_width = Inches(8.5)
    new_section.top_margin = Inches(0.75)
    new_section.bottom_margin = Inches(1.0)
    new_section.left_margin = Inches(0.625)
    new_section.right_margin = Inches(0.625)

    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)
    return new_section


def set_single_column_layout(doc):
    """Set layout to a single column (primarily for large images)."""
    new_section = doc.add_section()
    new_section.start_type = 0  # Continuous section break
    new_section.page_height = Inches(11)
    new_section.page_width = Inches(8.5)
    new_section.top_margin = Inches(0.75)
    new_section.bottom_margin = Inches(1.0)
    new_section.left_margin = Inches(0.625)
    new_section.right_margin = Inches(0.625)
    
    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)
    return new_section


def add_abstract_and_index_terms(doc, frontmatter):
    """Add the abstract and index terms block."""
    abstract_text = frontmatter.get('abstract', '')
    if abstract_text:
        abstract_para = doc.add_paragraph()
        run_bold = abstract_para.add_run('Abstract')
        run_bold.font.size = Pt(9)
        run_bold.font.bold = True
        run_bold.font.italic = True
        run_bold.font.name = 'Times New Roman'
        run_emdash = abstract_para.add_run('—')
        run_emdash.font.size = Pt(9)
        run_emdash.font.bold = True
        run_emdash.font.italic = True
        run_emdash.font.name = 'Times New Roman'
        run_text = abstract_para.add_run(abstract_text)
        run_text.font.size = Pt(9)
        run_text.font.bold = True
        run_text.font.name = 'Times New Roman'
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abstract_para.paragraph_format.space_after = Pt(9)

    index_terms_text = frontmatter.get('index_terms', '')
    if index_terms_text:
        index_para = doc.add_paragraph()
        run_bold = index_para.add_run('Index Terms')
        run_bold.font.size = Pt(9)
        run_bold.font.bold = True
        run_bold.font.italic = True
        run_bold.font.name = 'Times New Roman'
        run_emdash = index_para.add_run('—')
        run_emdash.font.size = Pt(9)
        run_emdash.font.bold = True
        run_emdash.font.italic = True
        run_emdash.font.name = 'Times New Roman'
        run_text = index_para.add_run(index_terms_text)
        run_text.font.size = Pt(9)
        run_text.font.name = 'Times New Roman'
        index_para.paragraph_format.space_after = Pt(12)


def process_table(doc, current_table):
    """Process a markdown table and append it correctly to the Word doc."""
    if len(current_table) > 0:
        rows = []
        for table_line in current_table:
            cells = [cell.strip() for cell in table_line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        if len(rows) > 0:
            max_cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(rows):
                row = table.rows[i]
                for j in range(min(len(row_data), max_cols)):
                    cell = row.cells[j]
                    cell.text = row_data[j].replace('**', '')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = 'Times New Roman'
                            if i == 0:
                                run.font.bold = True
                        if j == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)


def parse_markdown_content(doc, content, input_file):
    """Line by line parser to handle the body of the markdown."""
    lines = content.split('\n')
    skip_header = True
    in_code_block = False
    current_table = []
    in_table = False

    for line in lines:
        line_stripped = line.strip()
        
        # Skip until Introduction
        if 'I. INTRODUCTION' in line_stripped:
            skip_header = False
        if skip_header:
            continue
        
        # Main section headings
        if line_stripped.startswith('## I.') or line_stripped.startswith('## II.') or line_stripped.startswith('## III.') or line_stripped.startswith('## IV.') or line_stripped.startswith('## V.'):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped.replace('##', '').strip())
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(6)
        
        # Subsection headings
        elif line_stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped.replace('###', '').strip())
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.italic = True
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)

        # Images
        elif line_stripped.startswith('![') and '](' in line_stripped and line_stripped.endswith(')'):
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line_stripped)
            if match:
                img_path = match.group(2)
                
                # Resolve img_path relative to the markdown file directory
                md_dir = os.path.dirname(input_file)
                if md_dir:
                    img_path = os.path.join(md_dir, img_path)
                
                # Single column for large images
                set_single_column_layout(doc)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                if os.path.exists(img_path):
                    r = p.add_run()
                    try:
                        r.add_picture(img_path, width=Inches(7.25))
                        logger.info(f"Successfully inserted image: {img_path}")
                    except Exception as e:
                        logger.error(f"Failed to insert image {img_path}: {str(e)}")
                        p.add_run(f'[Error loading image: {img_path}]')
                else:
                    logger.warning(f"Image file not found: {img_path}")
                    p.add_run(f'[Image missing: {img_path}]')

        # Captions
        elif (line_stripped.startswith('*Fig. ') and line_stripped.endswith('*')) or (line_stripped.startswith('_Fig. ') and line_stripped.endswith('_')):
            caption_text = line_stripped.strip('*_')
            caption = doc.add_paragraph()
            run = caption.add_run(caption_text)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(12)
            
            # Resume two-column layout after caption
            set_two_column_layout(doc)
        
        # Code blocks
        elif line_stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        elif in_code_block:
            if line_stripped.startswith('CCS ='):
                p = doc.add_paragraph()
                run = p.add_run(line_stripped)
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            continue
        
        # Table detection
        elif line_stripped.startswith('|'):
            if all(c in '|-: ' for c in line_stripped):
                continue
            if not in_table:
                in_table = True
                current_table = []
            current_table.append(line_stripped)
        
        elif in_table and not line_stripped.startswith('|'):
            process_table(doc, current_table)
            in_table = False
            current_table = []
        
        # Acknowledgment section
        elif line_stripped.startswith('## ACKNOWLEDGMENT'):
            p = doc.add_paragraph()
            run = p.add_run('ACKNOWLEDGMENT')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(6)
        
        # References
        elif line_stripped.startswith('[') and ']' in line_stripped:
            if line_stripped.startswith('[1]'):
                ref_heading = doc.add_paragraph()
                run = ref_heading.add_run('REFERENCES')
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.name = 'Times New Roman'
                ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ref_heading.paragraph_format.space_before = Pt(9)
                ref_heading.paragraph_format.space_after = Pt(6)
            
            p = doc.add_paragraph()
            run = p.add_run(line_stripped)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(3)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Bullet points
        elif line_stripped.startswith('- ') or line_stripped.startswith('1. ') or line_stripped.startswith('2. ') or line_stripped.startswith('3. ') or line_stripped.startswith('4. ') or line_stripped.startswith('5. ') or line_stripped.startswith('6. '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped.replace('**', ''))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Regular paragraphs
        elif len(line_stripped) > 0 and not line_stripped.startswith('#') and not line_stripped.startswith('---') and not line_stripped.startswith('**Author'):
            cleaned_line = line_stripped.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(cleaned_line)
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0

    if in_table and current_table:
        process_table(doc, current_table)


def save_and_convert(doc, output_file):
    """Save DOCX and convert it to PDF."""
    try:
        doc.save(output_file)
        logger.info(f"IEEE Paper generated successfully in two-column format: {output_file}")
    except Exception as e:
        logger.error(f"Failed to save DOCX file {output_file}: {str(e)}\n{traceback.format_exc()}")
        sys.exit(1)

    try:
        logger.info("Attempting to convert DOCX to PDF using Word (docx2pdf)...")
        from docx2pdf import convert
        import os
        import subprocess
        import shutil
        
        abs_in = os.path.abspath(output_file)
        abs_out = os.path.abspath(output_file.replace('.docx', '.pdf'))
        
        # docx2pdf requires Word to be installed and not sandboxed
        try:
            convert(abs_in, abs_out)
            logger.info(f"IEEE Paper generated successfully in PDF format: {abs_out}")
            return
        except Exception as e:
            logger.warning(f"Word docx2pdf conversion failed. Falling back to alternative methods. Reason: {e}")
            
        # Fallback 1: LibreOffice Headless
        if shutil.which("soffice"):
            logger.info("Attempting fallback conversion using LibreOffice...")
            out_dir = os.path.dirname(abs_out)
            subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', out_dir, abs_in], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            logger.info(f"IEEE Paper generated successfully in PDF format using LibreOffice: {abs_out}")
            return
            
        # Fallback 2: Mac native textutil/cupsfilter routing (less layout perfect but works natively)
        logger.warning("No LibreOffice found. PDF must be manually generated from the provided DOCX.")
        logger.warning(f"Please open '{abs_in}' in Pages or Word and manually Export as PDF.")
        
    except ImportError:
        logger.error("docx2pdf module not found. Please run 'pip install docx2pdf' to enable PDF generation.")
    except Exception as e:
        logger.error(f"Failed to automatically convert DOCX to PDF: {str(e)}")


def generate_paper(input_file, output_file):
    """Main execution function to put it all together."""
    content = read_markdown_file(input_file)
    frontmatter, body_content = parse_frontmatter(content)
    
    doc = Document()
    setup_page_margins(doc)
    add_title(doc, frontmatter)
    add_authors(doc, frontmatter)
    
    set_two_column_layout(doc)
    add_abstract_and_index_terms(doc, frontmatter)
    
    # We now also pass `input_file` to `parse_markdown_content` so it can resolve images
    parse_markdown_content(doc, body_content, input_file)
    
    save_and_convert(doc, output_file)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate IEEE format paper from Markdown.')
    parser.add_argument('-i', '--input', type=str, default=DEFAULT_INPUT_FILE, help='Path to the input markdown file.')
    parser.add_argument('-o', '--output', type=str, default=DEFAULT_OUTPUT_FILE, help='Path to the output DOCX file.')
    
    args = parser.parse_args()
    
    logger.info(f"Starting generation with Input: {args.input} | Output: {args.output}")
    generate_paper(args.input, args.output)
